"""文档管理服务"""

from __future__ import annotations

import json
import logging
import os
from typing import Optional

logger = logging.getLogger(__name__)

from minio import Minio
from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.document import Document, Collection, User


class DocumentService:
    """文档管理服务"""

    def __init__(self):
        # 初始化 MinIO 客户端 (本地开发可无 MinIO)
        self.minio_available = False
        self.minio_client = None
        try:
            self.minio_client = Minio(
                endpoint=settings.MINIO_ENDPOINT,
                access_key=settings.MINIO_ACCESS_KEY,
                secret_key=settings.MINIO_SECRET_KEY,
                secure=settings.MINIO_USE_SSL,
            )
            self._ensure_bucket()
            self.minio_available = True
            logger.info(f"MinIO 连接成功: {settings.MINIO_ENDPOINT}")
        except Exception as e:
            logger.warning(f"MinIO 不可用 (本地开发模式): {e}")

    def _ensure_bucket(self):
        """确保存储桶存在"""
        if not self.minio_client:
            return
        try:
            if not self.minio_client.bucket_exists(settings.MINIO_BUCKET):
                self.minio_client.make_bucket(settings.MINIO_BUCKET)
        except Exception as e:
            logger.warning(f"MinIO bucket 检查失败: {e}")
            raise

    async def _get_storage_path(self, collection_id: str, filename: str) -> str:
        """获取 MinIO 存储路径"""
        import uuid
        unique_id = str(uuid.uuid4())
        ext = os.path.splitext(filename)[1]
        return f"{collection_id}/{unique_id}{ext}"

    async def create_collection(
        self,
        name: str,
        description: Optional[str] = None,
        owner_id: Optional[str] = None,
        chunk_size: int = 512,
        chunk_overlap: int = 128,
        db: Optional[AsyncSession] = None,
    ) -> Collection:
        """创建知识库集合"""
        from app.config import settings

        qdrant_collection = f"{settings.QDRANT_COLLECTION_PREFIX}{name.lower().replace(' ', '_')}"

        collection = Collection(
            name=name,
            description=description,
            qdrant_collection=qdrant_collection,
            owner_id=owner_id,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )

        if db:
            db.add(collection)
            await db.flush()
            await db.refresh(collection)

        return collection

    async def list_collections(
        self,
        db: AsyncSession,
        user: Optional[User] = None,
        skip: int = 0,
        limit: int = 100,
        tag_names: Optional[list[str]] = None,
    ) -> tuple[list[Collection], int]:
        """列出知识库集合。

        - 传 ``user``：admin 看全部；普通用户仅看自己有 ACL 的
        - 不传 ``user``：旧行为，返回全部（保留向后兼容）
        - ``tag_names``：按标签名称筛选（AND 逻辑）
        """
        from app.models.document import Tag, CollectionTag

        if user:
            from app.services.permission_service import PermissionService

            collections, total = await PermissionService().accessible_collections(user, db, skip, limit)
        else:
            total_result = await db.execute(select(func.count(Collection.id)))
            total = total_result.scalar() or 0

            result = await db.execute(
                select(Collection)
                .offset(skip)
                .limit(limit)
                .order_by(Collection.created_at.desc())
            )
            collections = list(result.scalars().all())

        # 标签筛选（AND 逻辑）
        if tag_names:
            filtered = []
            for c in collections:
                # 检查该知识库是否包含所有指定标签
                c_tag_names = {t.name.lower() for t in c.tags}
                if all(tn.lower() in c_tag_names for tn in tag_names):
                    filtered.append(c)
            collections = filtered
            total = len(collections)

        return collections, total

    async def upload_document(
        self,
        collection_id: str,
        filename: str,
        file_content: bytes,
        file_type: str,
        db: AsyncSession,
    ) -> Document:
        """上传文档"""
        # 存储文件到 MinIO (如果可用)
        if self.minio_available and self.minio_client:
            storage_path = await self._get_storage_path(collection_id, filename)
            try:
                self.minio_client.put_object(
                    bucket_name=settings.MINIO_BUCKET,
                    object_name=storage_path,
                    data=__import__("io").BytesIO(file_content),
                    length=len(file_content),
                    content_type=file_type,
                )
            except Exception as e:
                logger.warning(f"MinIO 存储失败 (使用本地路径): {e}")
                storage_path = f"local/{collection_id}/{filename}"
        else:
            storage_path = f"local/{collection_id}/{filename}"

        # 创建文档记录
        document = Document(
            collection_id=collection_id,
            filename=filename,
            file_path=storage_path,
            file_type=file_type,
            file_size=len(file_content),
            status="pending",
        )

        db.add(document)
        await db.flush()
        await db.refresh(document)

        return document

    async def list_documents(
        self,
        db: AsyncSession,
        user: Optional["User"] = None,
        collection_id: Optional[str] = None,
        skip: int = 0,
        limit: int = 100,
    ) -> tuple[list["Document"], int]:
        """列出文档

        - 传 ``user``：过滤为该用户可访问 KB 下的文档（admin 看过全部）
        - 传 ``collection_id``：仅返回该 KB 的文档（外层应已做权限校验）
        """
        from app.models.acl import CollectionACL

        base_query = select(Document)
        join_acl = False

        if user and user.role != "admin":
            # 普通用户：JOIN ACL 过滤
            base_query = base_query.join(
                CollectionACL,
                CollectionACL.collection_id == Document.collection_id,
            ).where(CollectionACL.user_id == user.id)
            join_acl = True

        if collection_id:
            base_query = base_query.where(Document.collection_id == collection_id)

        # count：用 func.count("*") 避免与 select_from 形成笛卡尔积
        count_target = base_query.subquery() if join_acl else base_query
        total = (
            await db.execute(
                select(func.count()).select_from(count_target)
            )
        ).scalar() or 0

        result = await db.execute(
            base_query.order_by(Document.created_at.desc())
            .offset(skip)
            .limit(limit)
        )
        return list(result.scalars().all()), total

    async def get_document(self, document_id: str, db: AsyncSession) -> Optional[Document]:
        """获取文档详情"""
        result = await db.execute(
            select(Document).where(Document.id == document_id)
        )
        return result.scalar_one_or_none()

    async def delete_document(self, document_id: str, db: AsyncSession) -> bool:
        """删除文档"""
        result = await db.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()
        if document is None:
            return False

        # 删除 MinIO 文件 (如果可用)
        if self.minio_available and self.minio_client:
            try:
                self.minio_client.remove_object(
                    bucket_name=settings.MINIO_BUCKET,
                    object_name=document.file_path,
                )
            except Exception:
                pass  # 文件可能已被删除

        await db.delete(document)
        return True

    async def download_file(self, document_id: str, db: AsyncSession) -> tuple[Optional[bytes], Optional[Document]]:
        """下载文件内容"""
        document = await self.get_document(document_id, db)
        if document is None:
            return None, None

        # 从 MinIO 读取
        if self.minio_available and self.minio_client:
            try:
                response = self.minio_client.get_object(
                    bucket_name=settings.MINIO_BUCKET,
                    object_name=document.file_path,
                )
                return response.read(), document
            except Exception as e:
                logger.warning(f"MinIO 读取失败: {e}")

        # 本地文件回退
        local_path = document.file_path.replace("local/", "")
        local_full_path = os.path.join(settings.UPLOAD_DIR or "uploads", local_path)
        if os.path.exists(local_full_path):
            with open(local_full_path, "rb") as f:
                return f.read(), document

        return None, document

    async def extract_text(self, document_id: str, db: AsyncSession) -> tuple[str, str]:
        """提取文档纯文本内容，返回 (content, format)"""
        from app.utils.file_utils import get_extension

        document = await self.get_document(document_id, db)
        if document is None:
            raise ValueError("文档不存在")

        file_content, doc = await self.download_file(document_id, db)
        if file_content is None:
            raise ValueError("文件内容不存在")

        ext = get_extension(doc.filename)

        try:
            if ext == ".pdf":
                from pypdf import PdfReader
                import io
                reader = PdfReader(io.BytesIO(file_content))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
                return text, "text"

            elif ext == ".docx":
                from docx import Document as DocxDocument
                import io
                docx = DocxDocument(io.BytesIO(file_content))
                text = "\n".join(p.text for p in docx.paragraphs if p.text)
                return text, "text"

            elif ext in (".md", ".txt", ".csv", ".json", ".xml", ".yaml", ".yml"):
                return file_content.decode("utf-8", errors="replace"), "text"

            elif ext in (".html", ".htm"):
                from html.parser import HTMLParser
                class TextExtractor(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.texts = []
                    def handle_data(self, data):
                        self.texts.append(data)
                parser = TextExtractor()
                parser.feed(file_content.decode("utf-8", errors="replace"))
                return "".join(parser.texts), "text"

            else:
                return "暂不支持该格式预览", "text"

        except Exception as e:
            logger.error(f"文本提取失败: {e}")
            return f"预览加载失败: {str(e)}", "text"
